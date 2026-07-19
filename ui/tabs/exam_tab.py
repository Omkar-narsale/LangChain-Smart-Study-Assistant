"""
ui/tabs/exam_tab.py — Bloom's Taxonomy Exam Generator
=====================================================
Premium glassmorphism UI for generating university-style exam questions.
Includes PDF and DOCX export.
"""

from __future__ import annotations
import io
import json
import logging
import time

import streamlit as st

from ui.components import render_empty_state
from prompts import exam_generator_prompt
from config import get_model, task_token_budget, invoke_with_retry
from langchain_core.output_parsers import StrOutputParser

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

logger = logging.getLogger(__name__)

# Single reusable parser
_parser = StrOutputParser()

# ---------------------------------------------------------------------------
# Caching Key
# ---------------------------------------------------------------------------
def _get_cache_key(marks: str, bloom: str, count: int) -> str:
    return f"{marks}_{bloom}_{count}"


# ---------------------------------------------------------------------------
# PDF / DOCX Generation
# ---------------------------------------------------------------------------
def _generate_pdf(questions: list[dict]) -> bytes:
    if not HAS_PDF:
        return b""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = styles["Heading1"]
    q_style = ParagraphStyle("Question", parent=styles["Normal"], fontName="Helvetica-Bold", spaceAfter=6)
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontName="Helvetica-Oblique", fontSize=9, textColor="#666666", spaceAfter=6)
    ans_style = ParagraphStyle("Answer", parent=styles["Normal"], fontName="Helvetica", spaceAfter=12)
    kw_style = ParagraphStyle("Keywords", parent=styles["Normal"], fontName="Helvetica-Oblique", fontSize=9, spaceAfter=20)
    
    story = []
    story.append(Paragraph("Smart Study - Exam Questions", title_style))
    story.append(Spacer(1, 12))
    
    for i, q in enumerate(questions, 1):
        story.append(Paragraph(f"Q{i}. {q.get('question', '')}", q_style))
        meta_text = f"Marks: {q.get('marks', '')} | Taxonomy: {q.get('taxonomy', '')} | Difficulty: {q.get('difficulty', '')}"
        story.append(Paragraph(meta_text, meta_style))
        story.append(Paragraph("Model Answer:", styles["Heading4"]))
        story.append(Paragraph(q.get('answer', ''), ans_style))
        kws = ", ".join(q.get("keywords", []))
        story.append(Paragraph(f"Keywords: {kws}", kw_style))
        story.append(Spacer(1, 12))
        
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def _generate_docx(questions: list[dict]) -> bytes:
    if not HAS_DOCX:
        return b""
    doc = docx.Document()
    doc.add_heading("Smart Study - Exam Questions", 0)
    
    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"Q{i}. {q.get('question', '')}", style="Heading 2")
        meta_text = f"Marks: {q.get('marks', '')} | Taxonomy: {q.get('taxonomy', '')} | Difficulty: {q.get('difficulty', '')}"
        doc.add_paragraph(meta_text, style="Intense Quote")
        
        doc.add_heading("Model Answer:", level=3)
        doc.add_paragraph(q.get("answer", ""))
        
        kws = ", ".join(q.get("keywords", []))
        doc.add_paragraph(f"Keywords: {kws}")
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ---------------------------------------------------------------------------
# LLM Invocation
# ---------------------------------------------------------------------------
def _generate_exam_questions(summary: str, key_points: str, marks: str, bloom_level: str, count: int) -> list[dict]:
    """Invoke the LLM to generate exam questions using only summary and key points."""
    with task_token_budget("exam_generator") as model:
        chain = exam_generator_prompt | model | _parser
        
        raw_output = invoke_with_retry(
            chain,
            {
                "summary": summary,
                "key_points": key_points,
                "marks": marks,
                "bloom_level": bloom_level,
                "count": count
            }
        )
        
    # Clean JSON
    cleaned = raw_output.strip()
    import re
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    
    try:
        data = json.loads(cleaned)
        return data.get("questions", [])
    except json.JSONDecodeError:
        # Fallback regex extraction if complete JSON fails
        match = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group())
                return data.get("questions", [])
            except:
                pass
                
        logger.error("Failed to parse JSON from exam generator: %s", raw_output)
        st.error("Failed to parse questions from AI. Please try again.")
        return []


# ---------------------------------------------------------------------------
# UI Renderer
# ---------------------------------------------------------------------------
def render_exam_tab(study_output: dict | None) -> None:
    """Render the Exam Generator tab."""
    st.markdown("<h2 class='text-section'>🎓 Bloom's Taxonomy Exam Generator</h2>", unsafe_allow_html=True)
    st.markdown("<p class='text-body'>Generate university-style questions and model answers.</p>", unsafe_allow_html=True)
    
    if not study_output:
        render_empty_state(
            "🎓",
            "No Study Material Generated",
            "Please generate study notes on the Study Notes tab first. The Exam Generator uses the summary and key points to minimize token usage."
        )
        return
        
    # Layout for inputs
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("<div style='font-size:13px;font-weight:700;color:var(--text-secondary);margin-bottom:8px;'>Section 1: Question Marks</div>", unsafe_allow_html=True)
        marks = st.radio(
            "Marks",
            ["2 Marks", "5 Marks", "10 Marks"],
            label_visibility="collapsed"
        )
        
    with col2:
        st.markdown("<div style='font-size:13px;font-weight:700;color:var(--text-secondary);margin-bottom:8px;'>Section 2: Bloom Level</div>", unsafe_allow_html=True)
        bloom = st.radio(
            "Bloom Level",
            ["Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create"],
            label_visibility="collapsed"
        )
        
    with col3:
        st.markdown("<div style='font-size:13px;font-weight:700;color:var(--text-secondary);margin-bottom:8px;'>Section 3: Number of Questions</div>", unsafe_allow_html=True)
        count_str = st.selectbox(
            "Number of Questions",
            ["1", "2", "3", "5", "10"],
            label_visibility="collapsed"
        )
        count = int(count_str)

    st.markdown("<br>", unsafe_allow_html=True)
    
    # Generate Button
    if st.button("🚀 Generate Questions", type="primary", use_container_width=True):
        summary = study_output.get("summary", "")
        # Convert key_points to string
        kp = study_output.get("key_points", [])
        kp_str = "\n".join(f"- {p}" for p in kp) if isinstance(kp, list) else str(kp)
        
        cache_key = _get_cache_key(marks, bloom, count)
        
        if cache_key not in st.session_state.exam_questions:
            with st.spinner(f"Generating {count} {bloom} questions for {marks}..."):
                questions = _generate_exam_questions(summary, kp_str, marks, bloom, count)
                if questions:
                    st.session_state.exam_questions[cache_key] = questions
                else:
                    return # Stop if generation failed
        else:
            # Inform user it was loaded from cache (optional)
            pass
            
        st.session_state.current_exam_key = cache_key
        st.rerun()

    # -----------------------------------------------------------------------
    # Display Results
    # -----------------------------------------------------------------------
    st.markdown("<hr style='border: 0; border-top: 1px solid var(--border); margin: 24px 0;'>", unsafe_allow_html=True)
    
    current_key = st.session_state.get("current_exam_key")
    if current_key and current_key in st.session_state.exam_questions:
        questions = st.session_state.exam_questions[current_key]
        
        st.markdown(f"### 📋 Generated Exam ({len(questions)} Questions)")
        
        # Action Buttons (Export)
        btn_col1, btn_col2, btn_col3 = st.columns([1, 1, 2])
        if HAS_PDF:
            with btn_col1:
                st.download_button(
                    "📄 Export PDF",
                    data=_generate_pdf(questions),
                    file_name="exam_questions.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        if HAS_DOCX:
            with btn_col2:
                st.download_button(
                    "📝 Export DOCX",
                    data=_generate_docx(questions),
                    file_name="exam_questions.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        for i, q in enumerate(questions):
            q_text = q.get("question", "")
            q_marks = q.get("marks", "")
            q_tax = q.get("taxonomy", "")
            q_diff = q.get("difficulty", "Medium")
            q_ans = q.get("answer", "")
            q_kws = ", ".join(q.get("keywords", []))
            
            # Badge Colors
            tax_color = "var(--primary)"
            diff_color = "#EAB308" if q_diff == "Medium" else ("#22C55E" if q_diff == "Easy" else "#EF4444")
            marks_color = "#3B82F6"
            
            st.markdown(f"""
            <div class="glass-card" style="margin-bottom: 20px; padding: 24px;">
                <div style="display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 16px;">
                    <div style="font-size: 18px; font-weight: 700; color: var(--text-primary); flex: 1; margin-right: 16px;">
                        <span style="color: var(--primary);">Q{i+1}.</span> {q_text}
                    </div>
                    <div style="display: flex; gap: 8px; flex-wrap: wrap; justify-content: flex-end;">
                        <span style="background: rgba(124,92,255,0.1); color: {tax_color}; border: 1px solid rgba(124,92,255,0.2); padding: 4px 10px; border-radius: 999px; font-size: 11px; font-weight: 700;">🧠 {q_tax}</span>
                        <span style="background: rgba(59,130,246,0.1); color: {marks_color}; border: 1px solid rgba(59,130,246,0.2); padding: 4px 10px; border-radius: 999px; font-size: 11px; font-weight: 700;">🎯 {q_marks}</span>
                        <span style="background: {diff_color}1a; color: {diff_color}; border: 1px solid {diff_color}33; padding: 4px 10px; border-radius: 999px; font-size: 11px; font-weight: 700;">📊 {q_diff}</span>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Interactive Expander for Answer
            with st.expander("👁️ Show Model Answer", expanded=False):
                st.markdown(f"""
                <div style="padding: 12px; background: rgba(0,0,0,0.2); border-radius: 8px; margin-bottom: 12px;">
                    <div style="font-size: 12px; font-weight: 700; color: var(--text-muted); text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 8px;">Model Answer</div>
                    <div style="font-size: 14px; color: var(--text-primary); line-height: 1.6;">{q_ans}</div>
                </div>
                <div style="padding: 12px; background: rgba(34,197,94,0.05); border-radius: 8px; border: 1px solid rgba(34,197,94,0.1);">
                    <div style="font-size: 12px; font-weight: 700; color: #4ade80; text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 4px;">Expected Keywords</div>
                    <div style="font-size: 13px; color: var(--text-secondary);">{q_kws}</div>
                </div>
                """, unsafe_allow_html=True)
