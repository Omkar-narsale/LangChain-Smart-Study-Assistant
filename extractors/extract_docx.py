"""
extract_docx.py — DOCX Text Extractor
=====================================
Extracts text from DOCX documents (paragraphs, headings, lists, tables).
"""

from __future__ import annotations
import io
import logging

import docx

logger = logging.getLogger(__name__)

def extract_docx(file: io.IOBase) -> str:
    """
    Extract text from a DOCX document, including tables.
    """
    try:
        document = docx.Document(file)
        content = []
        
        # Iterate over all elements (simplification: we'll extract paragraphs then tables)
        # Python-docx doesn't easily interleave paragraphs and tables in their document order 
        # without diving into oxml. We will extract paragraphs first, then tables for simplicity.
        # But wait, python-docx Document.iter_inner_content doesn't exist natively.
        # Let's extract paragraphs, then tables.
        
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                content.append(text)
                
        for table in document.tables:
            table_content = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_content.append(" | ".join(row_data))
            if table_content:
                content.append("Table Data:\n" + "\n".join(table_content))

        import streamlit as st
        st.session_state["page_texts"] = content
        return "\n\n".join(content)

    except Exception as exc:
        logger.exception("DOCX extraction failed.")
        raise RuntimeError(f"DOCX text extraction failed: {exc}") from exc
